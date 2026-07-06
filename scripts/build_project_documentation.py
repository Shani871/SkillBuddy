import ast
import re
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "SkillBuddy_Complete_Project_Documentation.docx"
ACCENT = "2E74B5"
DARK = "17365D"
MUTED = "667085"
LIGHT = "E8EEF5"
INK = "172B4D"

APP_DESCRIPTIONS = {
    "accounts": "Authentication, user roles, student/lecturer profiles, account administration, filters, and notifications.",
    "ai_tutor": "AI-assisted tutoring conversations, provider integration, safety checks, and chat persistence.",
    "chatbot": "General chatbot page and request handling.",
    "core": "Landing pages, news, academic sessions/semesters, dashboards, schedules, attendance, and calendars.",
    "course": "Programs, courses, allocations, enrollment, class schedules, academic events, and course content.",
    "emotions": "Student wellbeing check-ins, emotion analysis, alerts, dashboards, and resend workflows.",
    "payments": "Payment gateway pages and transaction workflows.",
    "quiz": "Quizzes, question types, sittings, marking, progress, and result presentation.",
    "result": "Assessment scores, grades, attendance summaries, transcripts, and result publication.",
    "search": "Cross-model search and template helpers.",
    "config": "Django settings, URL assembly, health checks, WSGI/ASGI, and Celery configuration.",
    "app": "Lightweight service entry points and API-facing helpers.",
    "scripts": "Administrative data-generation and setup utilities.",
}


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_w = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def add_page_field(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    run.font.color.rgb = RGBColor.from_string(MUTED)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)


def setup_styles(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = section.bottom_margin = Inches(0.8)
    section.left_margin = section.right_margin = Inches(0.85)
    section.header_distance = section.footer_distance = Inches(0.45)

    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.18
    for name, size, before, after, color in (
        ("Title", 28, 0, 8, DARK),
        ("Subtitle", 13, 0, 18, MUTED),
        ("Heading 1", 17, 18, 8, ACCENT),
        ("Heading 2", 13, 13, 6, ACCENT),
        ("Heading 3", 11, 9, 3, DARK),
    ):
        style = doc.styles[name]
        style.font.name = "Aptos Display" if name in ("Title", "Heading 1") else "Aptos"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = name != "Subtitle"
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    code = doc.styles.add_style("Code Metadata", WD_STYLE_TYPE.PARAGRAPH)
    code.font.name = "Consolas"
    code.font.size = Pt(8.5)
    code.font.color.rgb = RGBColor.from_string("344054")
    code.paragraph_format.space_after = Pt(3)
    code.paragraph_format.keep_together = True

    compact = doc.styles.add_style("Compact Body", WD_STYLE_TYPE.PARAGRAPH)
    compact.font.name = "Aptos"
    compact.font.size = Pt(9.5)
    compact.font.color.rgb = RGBColor.from_string(INK)
    compact.paragraph_format.space_after = Pt(4)
    compact.paragraph_format.line_spacing = 1.1

    header = section.header.paragraphs[0]
    header.text = "SKILLBUDDY  /  TECHNICAL REFERENCE"
    header.style = doc.styles["Caption"]
    header.runs[0].font.color.rgb = RGBColor.from_string(MUTED)
    header.runs[0].font.bold = True
    add_page_field(section.footer.paragraphs[0])


def add_kv_table(doc, rows):
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    for i, (label, value) in enumerate(rows):
        cells = table.rows[0].cells if i == 0 else table.add_row().cells
        cells[0].text = label
        cells[1].text = value
        set_cell_shading(cells[0], LIGHT)
        cells[0].paragraphs[0].runs[0].bold = True
    set_table_geometry(table, [2100, 7260])
    return table


def module_category(path):
    parts = path.parts
    if "migrations" in parts:
        return "Database migrations"
    if "tests" in parts or path.name == "tests.py":
        return "Automated tests"
    if parts[0] == "scripts":
        return "Administration scripts"
    return "Application code"


def signature(node):
    args = node.args
    items = []
    positional = list(args.posonlyargs) + list(args.args)
    defaults = [None] * (len(positional) - len(args.defaults)) + list(args.defaults)
    for arg, default in zip(positional, defaults):
        text = arg.arg
        if arg.annotation:
            text += ": " + ast.unparse(arg.annotation)
        if default is not None:
            text += " = " + ast.unparse(default)
        items.append(text)
    if args.vararg:
        items.append("*" + args.vararg.arg)
    elif args.kwonlyargs:
        items.append("*")
    for arg, default in zip(args.kwonlyargs, args.kw_defaults):
        text = arg.arg
        if default is not None:
            text += " = " + ast.unparse(default)
        items.append(text)
    if args.kwarg:
        items.append("**" + args.kwarg.arg)
    prefix = "async " if isinstance(node, ast.AsyncFunctionDef) else ""
    returns = f" -> {ast.unparse(node.returns)}" if node.returns else ""
    return f"{prefix}{node.name}({', '.join(items)}){returns}"


def humanize(name):
    return re.sub(r"_+", " ", name).strip().capitalize()


def infer_purpose(node, owner=None):
    doc = ast.get_docstring(node)
    if doc:
        return " ".join(doc.split())
    name = node.name
    patterns = (
        ("test_", "Verifies that "), ("get_", "Retrieves "), ("set_", "Sets "),
        ("create_", "Creates "), ("update_", "Updates "), ("delete_", "Deletes "),
        ("add_", "Adds "), ("edit_", "Edits "), ("validate_", "Validates "),
        ("clean_", "Validates and normalizes "), ("send_", "Sends "),
        ("generate_", "Generates "), ("calculate_", "Calculates "),
        ("mark_", "Marks "), ("is_", "Checks whether "), ("has_", "Checks whether "),
        ("can_", "Checks whether the caller can "), ("log_", "Records "),
    )
    if name == "__str__":
        return f"Returns the human-readable string representation of {owner or 'the object'}."
    if name == "__init__":
        return f"Initializes {owner or 'the object'} and its runtime state."
    if name == "save":
        return f"Persists {owner or 'the object'}, applying model-specific preprocessing."
    if name == "dispatch":
        return "Routes the request through class-based view dispatch logic."
    if name == "get_context_data":
        return "Builds template context for the class-based view."
    for prefix, verb in patterns:
        if name.startswith(prefix):
            return verb + humanize(name[len(prefix):]).lower() + "."
    calls_render = any(isinstance(x, ast.Call) and isinstance(x.func, ast.Name) and x.func.id == "render" for x in ast.walk(node))
    if calls_render:
        return f"Handles the {humanize(name).lower()} web request and renders its response."
    return f"Implements {humanize(name).lower()} behavior" + (f" for {owner}." if owner else ".")


def parse_modules():
    modules = []
    for path in sorted(ROOT.rglob("*.py")):
        rel = path.relative_to(ROOT)
        if rel == Path("scripts/build_project_documentation.py"):
            continue
        if any(part in {".git", ".venv", "venv", "staticfiles", "__pycache__"} for part in rel.parts):
            continue
        try:
            source = path.read_text(encoding="utf-8")
            tree = ast.parse(source)
        except (UnicodeDecodeError, SyntaxError):
            continue
        functions, classes = [], []

        def walk_body(body, owner=None):
            for node in body:
                if isinstance(node, ast.ClassDef):
                    bases = [ast.unparse(x) for x in node.bases]
                    methods = []
                    walk_body(node.body, node.name)
                    for child in node.body:
                        if isinstance(child, (ast.FunctionDef, ast.AsyncFunctionDef)):
                            methods.append(child)
                    classes.append((node, bases, methods))
                elif isinstance(node, (ast.FunctionDef, ast.AsyncFunctionDef)):
                    decorators = [ast.unparse(x) for x in node.decorator_list]
                    functions.append((node, owner, decorators))
                    walk_body(node.body, f"{owner + '.' if owner else ''}{node.name}")
                elif hasattr(node, "body") and isinstance(node.body, list):
                    walk_body(node.body, owner)

        walk_body(tree.body)
        if functions or classes:
            modules.append({
                "path": rel,
                "doc": ast.get_docstring(tree),
                "category": module_category(rel),
                "functions": functions,
                "classes": classes,
            })
    return modules


def add_function(doc, node, owner, decorators):
    qualified = f"{owner}.{node.name}" if owner else node.name
    p = doc.add_paragraph(style="Heading 3")
    p.add_run(qualified)
    meta = doc.add_paragraph(style="Code Metadata")
    meta.add_run(f"Line {node.lineno}  ·  {signature(node)}")
    if decorators:
        dec = doc.add_paragraph(style="Code Metadata")
        dec.add_run("Decorators: " + ", ".join("@" + x for x in decorators))
    purpose = doc.add_paragraph(style="Compact Body")
    purpose.add_run("Purpose: ").bold = True
    purpose.add_run(infer_purpose(node, owner))


def build_document():
    modules = parse_modules()
    function_count = sum(len(m["functions"]) for m in modules)
    class_count = sum(len(m["classes"]) for m in modules)
    doc = Document()
    setup_styles(doc)

    doc.add_paragraph("SKILLBUDDY", style="Title")
    doc.add_paragraph("Complete Project Documentation & Function Reference", style="Subtitle")
    p = doc.add_paragraph()
    p.add_run("A developer-oriented guide to the architecture, applications, data model, request flows, setup, and every Python function in the repository.").bold = True
    add_kv_table(doc, [
        ("Generated", date.today().strftime("%d %B %Y")),
        ("Repository", str(ROOT)),
        ("Coverage", f"{len(modules)} Python modules · {class_count} classes · {function_count} functions/methods"),
        ("Technology", "Django 4.2 · PostgreSQL/SQLite · Bootstrap · Celery/Redis · external AI and payment services"),
    ])
    doc.add_paragraph()
    p = doc.add_paragraph(style="Compact Body")
    p.add_run("Scope note: ").bold = True
    p.add_run("The function reference includes application code, model/form methods, utilities, scripts, migrations, and automated tests. Descriptions use source docstrings where available and concise behavior inference otherwise.")
    doc.add_page_break()

    doc.add_heading("1. Executive overview", level=1)
    doc.add_paragraph("SkillBuddy is a role-aware education management platform. It combines student records, courses and enrollment, schedules and attendance, assessments and quizzes, results, wellbeing monitoring, AI tutoring, search, news, and payment workflows in one Django project.")
    doc.add_heading("Primary user roles", level=2)
    for text in (
        "Administrators manage users, programs, courses, sessions, semesters, schedules, academic events, and platform reporting.",
        "Lecturers manage allocated courses, learning resources, assessments, scores, and student-facing academic activity.",
        "Students register for courses, view schedules and calendars, track attendance/results, take quizzes, and use tutoring and wellbeing features.",
        "Parents/guardians receive or review wellbeing-related information where enabled by the emotions application.",
    ):
        doc.add_paragraph(text, style="List Bullet")

    doc.add_heading("2. Architecture and request flow", level=1)
    doc.add_paragraph("The project follows Django’s model–template–view pattern. URL modules map browser requests to function-based or class-based views. Views apply authentication/role decorators, query models through Django ORM, and render HTML templates or return redirects/JSON/files. Forms perform validation; signals and Celery tasks handle side effects and asynchronous work.")
    add_kv_table(doc, [
        ("Presentation", "Django templates, Bootstrap 5, Font Awesome, static CSS/JavaScript"),
        ("Application", "Views, decorators, forms, filters, services, template tags"),
        ("Domain/data", "Django models for users, academics, assessments, results, alerts, and chat"),
        ("Infrastructure", "Django settings, PostgreSQL/SQLite, Redis/Celery, email, AI APIs, payment gateways"),
    ])

    doc.add_heading("3. Application map", level=1)
    app_counts = {}
    for module in modules:
        app = module["path"].parts[0]
        counts = app_counts.setdefault(app, [0, 0, 0])
        counts[0] += 1
        counts[1] += len(module["classes"])
        counts[2] += len(module["functions"])
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    for cell, text in zip(table.rows[0].cells, ("Area", "Responsibility", "Modules", "Classes / functions")):
        cell.text = text
        set_cell_shading(cell, ACCENT)
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        cell.paragraphs[0].runs[0].bold = True
    repeat_table_header(table.rows[0])
    for app, counts in sorted(app_counts.items()):
        row = table.add_row().cells
        row[0].text = app
        row[1].text = APP_DESCRIPTIONS.get(app, "Project support and implementation code.")
        row[2].text = str(counts[0])
        row[3].text = f"{counts[1]} / {counts[2]}"
    set_table_geometry(table, [1400, 5600, 900, 1460])

    doc.add_heading("4. Core data model", level=1)
    doc.add_paragraph("The principal relationships are centered on users, student profiles, programs/courses, enrollments, and academic activity. The exact field definitions and validation rules are authoritative in each app’s models.py file.")
    for title, detail in (
        ("Identity", "Custom User records carry role flags and connect to Student and lecturer-facing behavior."),
        ("Academic structure", "Programs contain Courses; Sessions and Semesters identify the active academic period."),
        ("Enrollment", "TakenCourse joins students to courses and stores grade-related state."),
        ("Planning", "ClassSchedule defines recurring teaching times; AcademicEvent represents dated exams, deadlines, holidays, and other events."),
        ("Assessment", "Quiz, question subclasses, sittings, progress, scores, and Result records represent evaluation workflows."),
        ("Engagement and wellbeing", "News/events, AI tutor messages, emotion check-ins, and alerts support student communication and assistance."),
    ):
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(title + ": ").bold = True
        p.add_run(detail)

    doc.add_heading("5. Security and role controls", level=1)
    doc.add_paragraph("Most protected views use Django login enforcement plus custom role decorators such as student_required, lecturer_required, or admin_required. CSRF middleware, Django password validation, environment-driven secrets, and ORM query parameterization provide the baseline security model. Production deployments should keep DEBUG disabled, use a strong SECRET_KEY, restrict ALLOWED_HOSTS, secure database/email/API credentials, and serve over HTTPS.")

    doc.add_heading("6. Local setup and operation", level=1)
    steps = (
        "Create and activate a Python virtual environment.",
        "Install requirements from requirements.txt (which selects the production dependency set) or the appropriate requirements file.",
        "Copy .env.example to .env and configure SECRET_KEY, DEBUG, ALLOWED_HOSTS, database, email, AI, payment, Redis, and Celery values as needed.",
        "Run python manage.py migrate to create/update the database schema.",
        "Create an administrator with python manage.py createsuperuser or the supplied helper script.",
        "Start the application with python manage.py runserver; start Redis/Celery separately when asynchronous features are enabled.",
        "Run the test suite with python manage.py test before deployment.",
    )
    for step in steps:
        doc.add_paragraph(step, style="List Number")

    doc.add_heading("7. Configuration reference", level=1)
    add_kv_table(doc, [
        ("Required baseline", "SECRET_KEY, DEBUG, ALLOWED_HOSTS, DJANGO_SETTINGS_MODULE, DATABASE_URL"),
        ("Email", "EMAIL_BACKEND, EMAIL_HOST, EMAIL_PORT, EMAIL_USE_TLS, EMAIL_HOST_USER, EMAIL_HOST_PASSWORD, EMAIL_FROM_ADDRESS"),
        ("AI", "GEMINI_API_KEY, GEMINI_MODEL"),
        ("Async", "USE_CELERY_EMOTION_ANALYSIS, USE_CELERY_EMAIL_NOTIFICATIONS, CELERY_BROKER_URL, CELERY_RESULT_BACKEND"),
        ("Payments", "STRIPE_SECRET_KEY, STRIPE_PUBLISHABLE_KEY, GOPAY_GOID, GOPAY_CLIENT_ID, GOPAY_CLIENT_SECRET, GOPAY_IS_PRODUCTION"),
        ("Identifiers", "STUDENT_ID_PREFIX, LECTURER_ID_PREFIX"),
    ])

    doc.add_heading("8. Testing and maintenance", level=1)
    doc.add_paragraph("Tests cover role restrictions, filtering, query behavior, email notifications, calendar visibility, attendance, quizzes, payments, AI tutoring, and emotion workflows. Migrations are included in the reference because they are executable project history, though they should normally be generated and applied through Django rather than edited manually.")

    doc.add_page_break()
    doc.add_heading("9. Complete Python API and function reference", level=1)
    doc.add_paragraph("Modules are grouped by project area and then by file. Each entry reports the source line and callable signature. Class methods are shown with their owning class; nested functions include their lexical owner.")

    current_app = None
    for module in modules:
        app = module["path"].parts[0]
        if app != current_app:
            doc.add_page_break()
            doc.add_heading(app.upper(), level=1)
            doc.add_paragraph(APP_DESCRIPTIONS.get(app, "Project implementation and support code."))
            current_app = app
        doc.add_heading(str(module["path"]), level=2)
        p = doc.add_paragraph(style="Compact Body")
        p.add_run(module["category"] + ". ").bold = True
        if module["doc"]:
            p.add_run(" ".join(module["doc"].split()))
        else:
            p.add_run(f"Contains {len(module['classes'])} class(es) and {len(module['functions'])} function/method definition(s).")
        if module["classes"]:
            doc.add_paragraph("Classes", style="Heading 3")
            for cls, bases, methods in module["classes"]:
                cp = doc.add_paragraph(style="Compact Body")
                cp.add_run(cls.name).bold = True
                cp.add_run(f" (line {cls.lineno})")
                if bases:
                    cp.add_run(" — extends " + ", ".join(bases))
                cdoc = ast.get_docstring(cls)
                if cdoc:
                    cp.add_run(". " + " ".join(cdoc.split()))
                else:
                    cp.add_run(f". Defines {len(methods)} direct method(s).")
        for node, owner, decorators in module["functions"]:
            add_function(doc, node, owner, decorators)

    doc.add_page_break()
    doc.add_heading("10. Handover checklist", level=1)
    for item in (
        "Confirm production environment variables and rotate any credentials that have ever been committed or shared.",
        "Run migrations and the complete automated test suite against the deployment database engine.",
        "Verify role-based access using representative admin, lecturer, student, and parent accounts.",
        "Validate email, AI provider, payment gateway, Redis, and Celery behavior in the target environment.",
        "Review accessibility, localization catalogs, backups, logging, monitoring, and incident-response procedures.",
        "Update this document after material schema, route, integration, or workflow changes.",
    ):
        doc.add_paragraph(item, style="List Bullet")

    props = doc.core_properties
    props.title = "SkillBuddy Complete Project Documentation"
    props.subject = "Architecture, setup, modules, classes, and complete Python function reference"
    props.author = "SkillBuddy Engineering"
    props.keywords = "SkillBuddy, Django, technical documentation, function reference"
    doc.save(OUT)
    print(f"Wrote {OUT} ({len(modules)} modules, {class_count} classes, {function_count} functions)")


if __name__ == "__main__":
    build_document()
